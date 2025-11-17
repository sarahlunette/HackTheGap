"""
Qdrant Vectorstore Builder (v6)
- pdfplumber for PDF extraction
- DOCX, TXT, MD, CSV support
- NEW: JSON ingestion (arrays, dicts, nested)
- LOCAL embedding model: ../models/all-MiniLM-L6-v2
- Automatic sync with Qdrant:
      ‚úî add new files
      ‚úî remove deleted files
"""

from pathlib import Path
import pandas as pd
import pdfplumber
import json

from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct

from llama_index.core import Document
from llama_index.embeddings.huggingface import HuggingFaceEmbedding



# ============================================================
# CONFIG
# ============================================================
DOCS_DIR = Path("./docs")

QDRANT_URL = "http://localhost:6333"
COLLECTION_NAME = "island_docs"

EMBEDDING_MODEL = "./models/all-MiniLM-L6-v2"



# ============================================================
# HELPERS ‚Äî Loaders
# ============================================================
def load_pdf(path: Path) -> list[Document]:
    docs = []
    try:
        with pdfplumber.open(path) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)

        if text.strip():
            docs.append(Document(text=text, metadata={"source": path.name}))
        else:
            print(f"‚ö† PDF EMPTY: {path.name}")

    except Exception as e:
        print(f"‚ùå Error reading PDF {path.name}: {e}")

    return docs


def load_csv_rows(path: Path) -> list[Document]:
    df = pd.read_csv(path)
    docs = []

    for i, row in df.iterrows():
        text = "\n".join(f"{col}: {row[col]}" for col in df.columns)

        if text.strip():
            docs.append(
                Document(
                    text=text,
                    metadata={"source": path.name, "row_index": int(i)}
                )
            )

    return docs


def load_text_file(path: Path) -> list[Document]:
    try:
        text = path.read_text(errors="ignore")
        if text.strip():
            return [Document(text=text, metadata={"source": path.name})]
    except:
        pass
    return []


def load_docx(path: Path) -> list[Document]:
    try:
        import docx
        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return [Document(text=text, metadata={"source": path.name})]
    except:
        print(f"‚ùå Cannot read DOCX: {path.name}")
        return []


# ============================================================
# NEW ‚Äî JSON ingestion
# ============================================================
def load_json(path: Path) -> list[Document]:
    """
    Loads JSON files with the following rules:
    - If JSON is a list ‚Üí each element becomes one Document
    - If JSON is a dict ‚Üí the dict becomes one Document
    - Nested structures are flattened into readable text
    """
    try:
        data = json.loads(path.read_text(errors="ignore"))
    except Exception as e:
        print(f"‚ùå Error reading JSON {path.name}: {e}")
        return []

    docs = []

    def flatten(obj, prefix=""):
        """Recursively flatten a dict/list into readable text."""
        lines = []

        if isinstance(obj, dict):
            for k, v in obj.items():
                key = f"{prefix}{k}"
                if isinstance(v, (dict, list)):
                    lines.append(f"{key}:")
                    lines.append(flatten(v, prefix=prefix + "  "))
                else:
                    lines.append(f"{key}: {v}")

        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                lines.append(f"{prefix}- item {i}:")
                lines.append(flatten(item, prefix=prefix + "  "))

        else:
            lines.append(f"{prefix}{obj}")

        return "\n".join(lines)

    # Case 1: List of JSON objects
    if isinstance(data, list):
        for idx, item in enumerate(data):
            text = flatten(item)
            docs.append(
                Document(
                    text=text,
                    metadata={"source": path.name, "json_index": idx}
                )
            )

    # Case 2: Single dict
    elif isinstance(data, dict):
        text = flatten(data)
        docs.append(Document(text=text, metadata={"source": path.name}))

    else:
        print(f"‚ö† Unsupported JSON structure in {path.name}")

    return docs



# ============================================================
# KEY GENERATOR FOR DOC IDENTIFICATION
# ============================================================
def doc_key(doc: Document) -> str:
    """Unique ID for each document based on source + index for CSV/JSON rows."""
    key = doc.metadata.get("source")

    for suffix in ("row_index", "json_index"):
        if suffix in doc.metadata:
            key += f"#{suffix}:{doc.metadata[suffix]}"

    return key



# ============================================================
# SYNC FUNCTION ‚Äî Add new docs / remove deleted ones
# ============================================================
def sync_vectorstore(client, collection_name, embed_model, docs):
    print("\nüîÑ Syncing Qdrant collection‚Ä¶")

    try:
        existing, _ = client.scroll(collection_name, limit=50_000)
    except Exception as e:
        print(f"‚ùå Could not scroll existing vectors: {e}")
        return

    existing_index = {}
    for point in existing:
        key = point.payload.get("source", "")
        for suffix in ("row_index", "json_index"):
            if suffix in point.payload:
                key += f"#{suffix}:{point.payload[suffix]}"
        existing_index[key] = point.id

    new_index = {doc_key(d): d for d in docs}

    # ---------------
    # Detect deletions
    # ---------------
    to_delete = [existing_index[k] for k in existing_index if k not in new_index]

    if to_delete:
        print(f"üóë Removing {len(to_delete)} removed docs‚Ä¶")
        client.delete(collection_name, points_selector={"points": to_delete})
    else:
        print("üü¢ No deletions needed.")

    # ---------------
    # Detect additions
    # ---------------
    to_add = [new_index[k] for k in new_index if k not in existing_index]

    if to_add:
        print(f"‚ûï Adding {len(to_add)} new docs‚Ä¶")

        next_id = max(existing_index.values(), default=0) + 1
        new_points = []

        for doc in to_add:
            vector = embed_model.get_text_embedding(doc.text)
            new_points.append(
                PointStruct(
                    id=next_id,
                    vector=vector,
                    payload={"text": doc.text, **doc.metadata}
                )
            )
            next_id += 1

        client.upsert(collection_name, points=new_points)
    else:
        print("üü¢ No new additions.")

    print("‚úî Sync complete.")



# ============================================================
# MAIN
# ============================================================
def main():
    print("\n=== üöÄ QDRANT VECTORSTORE BUILDER v6 ===")

    if not DOCS_DIR.exists():
        raise FileNotFoundError("docs/ folder missing")

    # ------------------------------------------------------------
    # Load documents
    # ------------------------------------------------------------
    all_docs = []
    print("üìö Loading documents‚Ä¶")

    for file in DOCS_DIR.iterdir():
        suffix = file.suffix.lower()

        if suffix == ".pdf":
            print(f"  ‚Üí PDF: {file.name}")
            all_docs.extend(load_pdf(file))

        elif suffix in (".txt", ".md"):
            print(f"  ‚Üí TEXT: {file.name}")
            all_docs.extend(load_text_file(file))

        elif suffix == ".docx":
            print(f"  ‚Üí DOCX: {file.name}")
            all_docs.extend(load_docx(file))

        elif suffix == ".csv":
            print(f"  ‚Üí CSV rows: {file.name}")
            all_docs.extend(load_csv_rows(file))

        elif suffix == ".json":
            print(f"  ‚Üí JSON: {file.name}")
            all_docs.extend(load_json(file))

    print(f"\nüì¶ Loaded documents: {len(all_docs)}")

    if not all_docs:
        print("‚ö† No documents ‚Äî stopping.")
        return

    # ------------------------------------------------------------
    # Embedding model
    # ------------------------------------------------------------
    print(f"\nüß† Loading embedding model: {EMBEDDING_MODEL}")
    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)

    embedding_dim = len(embed_model.get_text_embedding("test"))
    print(f"üî¢ Embedding size = {embedding_dim}")

    # ------------------------------------------------------------
    # Qdrant connection
    # ------------------------------------------------------------
    print("\nüóÑ Connecting to Qdrant‚Ä¶")
    client = QdrantClient(url=QDRANT_URL)

    collections = [c.name for c in client.get_collections().collections]

    # ------------------------------------------------------------
    # Create or Sync
    # ------------------------------------------------------------
    if COLLECTION_NAME not in collections:
        print(f"üìå Creating new collection `{COLLECTION_NAME}`")

        client.recreate_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=embedding_dim, distance=Distance.COSINE),
        )

        points = []
        pid = 1

        for doc in all_docs:
            vector = embed_model.get_text_embedding(doc.text)
            points.append(
                PointStruct(
                    id=pid,
                    vector=vector,
                    payload={"text": doc.text, **doc.metadata}
                )
            )
            pid += 1

        client.upsert(collection_name=COLLECTION_NAME, points=points)
        print("‚úî Full ingestion complete.")

    else:
        print(f"üîÑ Collection exists ‚Äî syncing‚Ä¶")
        sync_vectorstore(client, COLLECTION_NAME, embed_model, all_docs)

    # ------------------------------------------------------------
    # Stats
    # ------------------------------------------------------------
    count = client.count(COLLECTION_NAME).count
    print("\nüéâ DONE ‚Äî Qdrant Vectorstore Ready!")
    print(f"üìå Collection: {COLLECTION_NAME}")
    print(f"üìä Total vectors stored: {count}")
    print("üöÄ Ready for RAG pipelines.\n")



if __name__ == "__main__":
    main()
